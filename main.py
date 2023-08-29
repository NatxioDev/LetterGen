import os
from datetime import datetime
import pandas as pd
from docx import Document


# Parameters
TEMPLATE_LETTER = './templates/sample.docx'
DATA_CSV = './data/info.csv'
OUTPUT_FOLDER = './output'


def verify_folder():
    if not os.path.exists('./output'):
        os.makedirs('./output')
        print("Folder 'output' was not found, created")
    else:
        print("Folder 'output' already exists")

    if not os.path.exists('./templates'):
        print("Folder 'templates' was not found, please create it and put the template letter inside")
        exit()

    if not os.path.exists('./data'):
        print("Folder 'data' was not found, please create it and put the csv file inside")
        exit()

def current_date_format(date):
    months = ("enero", "febrero", "marzo", "abril", "mayo", "junio", "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre")
    day = date.day
    month = months[date.month - 1]
    year = date.year
    messsage = "{} de {}, {}".format(day, month, year)

    return messsage

def clean_folder():
    for file in os.listdir(OUTPUT_FOLDER):
        os.remove(OUTPUT_FOLDER + '/{}'.format(file))
    
def indentify_template():
    # Print paragraphs in template with index

    template = Document(TEMPLATE_LETTER)

    for i in range(len(template.paragraphs)):
        print(i, template.paragraphs[i].text)

def main():
# Get template letter
    template = Document(TEMPLATE_LETTER)

    # Get data from csv and iterate over it
    data = pd.read_csv(DATA_CSV, sep=',')

    for index, row in data.iterrows():
        if not pd.isna(row['Ponente']):
            
            date = template.paragraphs[2]
            srsra = template.paragraphs[4]
            name = template.paragraphs[5]
            deno = template.paragraphs[6]

            # Get today date and replace in template
            now = datetime.now()
            date.text = "La Paz, " + current_date_format(now) 

            # Get srsra and replace in template
            # srsra.text = "??"
    
            # Get name and replace in template
            name.text = row['Ponente']

            if not pd.isna(row['Empresa']) and not pd.isna(row['Cargo']):
                deno.text = row['Cargo'] + " - " + row['Empresa']

            elif not pd.isna(row['Empresa']):
                deno.text = row['Empresa']

            elif not pd.isna(row['Cargo']):
                deno.text = row['Cargo']
            
            else:
                deno.text = ""

            
            # Save letter in output folder using const OUTPUT_FOLDER
            template.save(OUTPUT_FOLDER + '/{}.docx'.format(row['Ponente']))


            template.save('./output/{}.docx'.format(row['Ponente']))
            print("Nro. {} -> Letter {} created".format(index, row['Ponente']))


# Use this to indentify the index of each paragraph in the template ⬇️
# indentify_template()

# Run the program 
# You need to create the folders 'templates' and 'data' and put the files inside them 
# Also you would change the parameters or the filenames in the top of the file
verify_folder()
clean_folder()
main()